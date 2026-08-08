from __future__ import annotations

import io
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402

from app.article_reader import (  # noqa: E402
    ARTICLE_EXPORT_FORMATS,
    article_to_docx,
    parse_wechat_article_html,
    render_article_export,
    write_article_export,
)

SAMPLE = """
<html><body>
<div id="js_article">
  <h1 id="activity-name">Word导出测试</h1>
  <div id="js_content"><p>段落一。</p><p>段落二。</p></div>
</div>
</body></html>
"""

RICH_SAMPLE = """
<html><body>
<div id="js_article">
  <h1 id="activity-name">富文本测试</h1>
  <div id="js_content">
    <p>含<strong>加粗</strong>与<em>斜体</em>。</p>
    <h2>小节</h2>
    <ul><li>甲</li><li>乙</li></ul>
    <ol><li>一</li><li>二</li></ol>
    <blockquote>引用一句。</blockquote>
    <p><img src="https://invalid.example.test/no_such_img.png" /></p>
  </div>
</div>
</body></html>
"""


def test_word_in_export_formats():
    assert "word" in ARTICLE_EXPORT_FORMATS
    assert ARTICLE_EXPORT_FORMATS["word"] in ("Word", "Word (.docx)")


def test_article_to_docx_is_zip_container():
    art = parse_wechat_article_html(SAMPLE, source_url="https://mp.weixin.qq.com/s/abc")
    data = article_to_docx(art)
    assert isinstance(data, (bytes, bytearray))
    assert data[:2] == b"PK"


def test_write_article_export_docx(tmp_path: Path | None = None):
    base = tmp_path if tmp_path is not None else Path(__file__).resolve().parents[1] / "data" / "_test_exports"
    base.mkdir(parents=True, exist_ok=True)
    art = parse_wechat_article_html(SAMPLE, source_url="https://mp.weixin.qq.com/s/abc")
    out = base / "word_test.docx"
    if out.exists():
        out.unlink()
    path = write_article_export(out, art, "word")
    assert path.is_file()
    assert path.read_bytes()[:2] == b"PK"
    # text formats still via render
    md = render_article_export(art, "markdown")
    assert md.startswith("# ")


def _parse_docx(art: dict) -> Document:
    return Document(io.BytesIO(article_to_docx(art)))


def test_docx_rich_text_structure():
    art = parse_wechat_article_html(RICH_SAMPLE, source_url="https://mp.weixin.qq.com/s/rich")
    doc = _parse_docx(art)
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)
    assert "富文本测试" in joined
    assert "加粗" in joined and "斜体" in joined
    assert "小节" in joined
    assert "甲" in joined and "乙" in joined
    assert "一" in joined and "二" in joined
    assert "引用一句" in joined
    # 标题用 Title 样式
    assert any(p.style.name == "Title" and "富文本测试" in p.text for p in doc.paragraphs)
    # 小节用 Heading 样式
    assert any(p.style.name.startswith("Heading") and "小节" in p.text for p in doc.paragraphs)
    # 图片无法下载时保留占位文本
    assert "[图片]" in joined


def test_docx_list_numbering_sequence():
    art = parse_wechat_article_html(RICH_SAMPLE, source_url="https://mp.weixin.qq.com/s/rich")
    doc = _parse_docx(art)
    ordered = [p.text for p in doc.paragraphs if p.text.strip().startswith(("1. ", "2. "))]
    assert ordered == ["1. 一", "2. 二"]


def test_docx_hyperlink_run_present():
    art = parse_wechat_article_html(
        '<div id="js_article"><h1 id="activity-name">链</h1>'
        '<div id="js_content"><p>看<a href="https://example.com">这里</a></p></div></div>',
        source_url="https://mp.weixin.qq.com/s/link",
    )
    doc = _parse_docx(art)
    body = doc.paragraphs[1]._element.xml
    assert "w:hyperlink" in body
    # 超链接 URL 存放在 part.rels（OOXML relationship），不在 document.xml
    rel_targets = [r.target_ref for r in doc.part.rels.values()]
    assert "https://example.com" in rel_targets
