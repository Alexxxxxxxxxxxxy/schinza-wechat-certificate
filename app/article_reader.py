"""Read WeChat MP article HTML and export (inspired by wechat-article-exporter).

Formats: html (normalized WeChat layout), markdown, txt, json, word (.docx).
"""

from __future__ import annotations

from app.errors import describe_exception

import csv
import html
import io
import json
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Callable
from urllib.parse import unquote

import requests
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36 "
    "MicroMessenger/7.0.20.1781(0x6700143B) NetType/WIFI "
    "WindowsWechat(0x63090a13) XWEB/11275"
)

ARTICLE_EXPORT_FORMATS: dict[str, str] = {
    "html": "HTML",
    "markdown": "Markdown",
    "txt": "TXT",
    "json": "JSON",
    "csv": "CSV",
    "word": "Word",
}

ARTICLE_EXPORT_LABELS = list(ARTICLE_EXPORT_FORMATS.values())


def _fully_unquote(value: str) -> str:
    s = value or ""
    for _ in range(3):
        n = unquote(s)
        if n == s:
            break
        s = n
    return s


def _html_to_text(fragment: str) -> str:
    if not fragment:
        return ""
    soup = BeautifulSoup(fragment, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    text = soup.get_text("\n", strip=True)
    lines = [ln.strip() for ln in text.splitlines()]
    return "\n".join(ln for ln in lines if ln)


def _extract_publish_ts(html_text: str) -> int:
    for pattern in (
        r'var\s+ct\s*=\s*"(\d+)"',
        r'var\s+createTime\s*=\s*[\'"](\d+)[\'"]',
        r'publish_time\s*[:=]\s*[\'"]?(\d{10})',
        r'content_noencode.*?createTime\s*[:=]\s*[\'"]?(\d{10})',
    ):
        m = re.search(pattern, html_text or "", re.I | re.S)
        if m:
            try:
                ts = int(m.group(1))
                if ts > 1_000_000_000:
                    return ts
            except Exception:
                continue
    return 0


def normalize_wechat_html(raw_html: str) -> str:
    """Normalize raw MP HTML like wechat-article-exporter's normalizeHtml.

    - Prefer #js_article shell when present
    - Unhide #js_content
    - Drop ads / QR / scripts
    - Promote data-src images to src
    """
    soup = BeautifulSoup(raw_html or "", "html.parser")
    article = soup.select_one("#js_article")
    if article is None:
        # Fallback: wrap parsed body content
        content = soup.select_one("#js_content") or soup.select_one("div.rich_media_content")
        title_el = soup.select_one("#activity-name") or soup.select_one("h1.rich_media_title")
        title = title_el.get_text(strip=True) if title_el else ""
        body_inner = str(content) if content else _html_to_text(raw_html or "")
        return article_to_html_document(
            {
                "title": title or "(无标题)",
                "body_html": body_inner,
                "body_text": _html_to_text(body_inner),
                "link": "",
                "publish_at": "",
            }
        )

    content = article.select_one("#js_content")
    if content is not None and isinstance(content, Tag):
        if content.has_attr("style"):
            del content["style"]

    for sel in (
        "#js_top_ad_area",
        "#js_tags_preview_toast",
        "#content_bottom_area",
        "#js_pc_qr_code",
        "#wx_stream_article_slide_tip",
        "script",
    ):
        for node in article.select(sel):
            node.decompose()

    for img in soup.find_all("img"):
        if not isinstance(img, Tag):
            continue
        src = img.get("src") or img.get("data-src")
        if src:
            img["src"] = src

    body = soup.body
    body_cls = ""
    if body is not None and isinstance(body, Tag):
        body_cls = " ".join(body.get("class") or [])

    page = str(article)
    return (
        "<!DOCTYPE html>\n"
        '<html lang="zh_CN">\n'
        "<head>\n"
        '<meta charset="utf-8">\n'
        '<meta name="viewport" content="width=device-width,initial-scale=1.0">\n'
        '<meta name="referrer" content="no-referrer">\n'
        "<style>\n"
        "#page-content,#js_article,.__page_content__{max-width:667px;margin:0 auto;}\n"
        "img{max-width:100%;height:auto;}\n"
        "body{font-family:Microsoft YaHei,Segoe UI,sans-serif;line-height:1.7;"
        "color:#1a1a1a;background:#fff;padding:16px;}\n"
        "</style>\n"
        "</head>\n"
        f'<body class="{html.escape(body_cls)}">\n'
        f"{page}\n"
        "</body>\n"
        "</html>\n"
    )


def parse_wechat_article_html(
    html_text: str,
    *,
    source_url: str = "",
) -> dict[str, Any]:
    soup = BeautifulSoup(html_text or "", "html.parser")
    title = ""
    title_el = soup.select_one("#activity-name") or soup.select_one("h1.rich_media_title")
    if title_el:
        title = title_el.get_text(strip=True)
    og_title = soup.find("meta", property="og:title")
    if og_title and og_title.get("content"):
        title = title or str(og_title["content"]).strip()

    content = soup.select_one("#js_content") or soup.select_one("div.rich_media_content")
    body_html = str(content) if content else ""
    body_text = _html_to_text(body_html) if body_html else _html_to_text(html_text or "")

    if len(body_text) < 20:
        og_desc = soup.find("meta", property="og:description")
        if og_desc and og_desc.get("content"):
            desc = str(og_desc["content"]).strip()
            if len(desc) > len(body_text):
                body_text = desc

    publish_ts = _extract_publish_ts(html_text or "")
    publish_at = (
        datetime.fromtimestamp(publish_ts).strftime("%Y-%m-%d %H:%M")
        if publish_ts
        else ""
    )

    normalized = ""
    try:
        if soup.select_one("#js_article") or content:
            normalized = normalize_wechat_html(html_text or "")
    except Exception:
        normalized = ""

    return {
        "title": title or "(无标题)",
        "body_text": body_text,
        "body_html": body_html or body_text,
        "normalized_html": normalized,
        "link": source_url or "",
        "publish_ts": publish_ts,
        "publish_at": publish_at,
        "videos": extract_videos_from_html(body_html or html_text),
    }


def article_to_markdown(art: dict[str, Any]) -> str:
    title = str(art.get("title") or "(无标题)").strip()
    link = str(art.get("link") or "").strip()
    when = str(art.get("publish_at") or "").strip()
    body_html = str(art.get("body_html") or "")
    body = _html_fragment_to_markdown(body_html) if body_html else ""
    if not body.strip():
        body = str(art.get("body_text") or "").strip()
    lines = [f"# {title}", ""]
    if link:
        lines.append(f"来源：{link}")
    if when:
        lines.append(f"发布时间：{when}")
    if link or when:
        lines.append("")
    lines.append(body)
    lines.append("")
    lines.extend(_video_lines(art))
    return "\n".join(lines)


def _video_lines(art: dict[str, Any]) -> list[str]:
    """Return text lines listing an article's videos (downloaded path or URL)."""
    videos = [v for v in (art.get("videos") or []) if isinstance(v, dict)]
    if not videos:
        return []
    lines = ["", "视频："]
    for v in videos:
        target = str(v.get("local_path") or v.get("url") or "").strip()
        if target:
            lines.append(f"  {target}")
    return lines


def _html_fragment_to_markdown(fragment: str) -> str:
    """Lightweight HTML→MD (headings, paragraphs, lists, links, images, bold/italic)."""
    if not fragment:
        return ""
    soup = BeautifulSoup(fragment, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()

    def walk(node: Any) -> str:
        if isinstance(node, str):
            return re.sub(r"\s+", " ", node)
        if not isinstance(node, Tag):
            return ""
        name = (node.name or "").lower()
        inner = "".join(walk(c) for c in node.children)
        if name in ("h1", "h2", "h3", "h4"):
            level = int(name[1])
            return f"\n{'#' * level} {inner.strip()}\n\n"
        if name == "p":
            return f"\n{inner.strip()}\n\n"
        if name == "br":
            return "\n"
        if name in ("strong", "b"):
            return f"**{inner.strip()}**" if inner.strip() else ""
        if name in ("em", "i"):
            return f"*{inner.strip()}*" if inner.strip() else ""
        if name == "a":
            href = node.get("href") or ""
            text = inner.strip() or href
            return f"[{text}]({href})" if href else text
        if name == "img":
            src = node.get("src") or node.get("data-src") or ""
            alt = node.get("alt") or "image"
            return f"\n![{alt}]({src})\n\n" if src else ""
        if name in ("ul", "ol"):
            items = []
            for i, li in enumerate(node.find_all("li", recursive=False), start=1):
                t = "".join(walk(c) for c in li.children).strip()
                prefix = f"{i}." if name == "ol" else "-"
                items.append(f"{prefix} {t}")
            return "\n" + "\n".join(items) + "\n\n"
        if name == "blockquote":
            quoted = "\n".join("> " + ln for ln in inner.strip().splitlines() if ln.strip())
            return f"\n{quoted}\n\n"
        if name in ("div", "section", "span", "section"):
            return inner
        return inner

    text = walk(soup)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


def article_to_txt(art: dict[str, Any]) -> str:
    title = str(art.get("title") or "(无标题)").strip()
    link = str(art.get("link") or "").strip()
    when = str(art.get("publish_at") or "").strip()
    body = str(art.get("body_text") or "").strip()
    lines = [title, ""]
    if link:
        lines.append(f"来源：{link}")
    if when:
        lines.append(f"发布时间：{when}")
    if link or when:
        lines.append("")
    lines.append(body)
    lines.append("")
    lines.extend(_video_lines(art))
    return "\n".join(lines)


def article_to_json(art: dict[str, Any]) -> str:
    payload = {
        "title": art.get("title") or "(无标题)",
        "link": art.get("link") or "",
        "publish_ts": int(art.get("publish_ts") or 0),
        "publish_at": art.get("publish_at") or "",
        "body_text": art.get("body_text") or "",
        "body_html": art.get("body_html") or "",
        "videos": art.get("videos") or [],
        "exported_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


CSV_COLUMNS = ("标题", "链接", "发布时间", "作者", "摘要", "正文", "视频")


def article_to_csv_row(art: dict[str, Any]) -> dict[str, str]:
    """One CSV row for an article (标题/链接/时间/作者/摘要/正文)."""
    return {
        "标题": str(art.get("title") or "(无标题)"),
        "链接": str(art.get("link") or ""),
        "发布时间": str(art.get("publish_at") or ""),
        "作者": str(art.get("author") or ""),
        "摘要": str(art.get("digest") or ""),
        "正文": str(art.get("body_text") or "").strip(),
        "视频": " | ".join(
            str(v.get("local_path") or v.get("url") or "")
            for v in (art.get("videos") or [])
            if isinstance(v, dict) and (v.get("local_path") or v.get("url"))
        ),
    }


def _write_csv(path: Path, rows: list[dict[str, str]]) -> None:
    """Write CSV with utf-8-sig BOM so Excel opens Chinese correctly."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=list(CSV_COLUMNS))
        writer.writeheader()
        writer.writerows(rows)


VIDEO_DL_DIRNAME = "视频"


def extract_videos_from_html(raw_html: str) -> list[dict[str, str]]:
    """Extract videos from article HTML: WeChat CDN mp4 + Tencent v.qq.com iframes."""
    videos: list[dict[str, str]] = []
    if not raw_html:
        return videos
    seen: set[str] = set()
    try:
        soup = BeautifulSoup(raw_html or "", "html.parser")
    except Exception:
        return videos

    def add(kind: str, url: str, vid: str = "") -> None:
        url = html.unescape((url or "").strip())
        if not url or url in seen:
            return
        seen.add(url)
        item: dict[str, str] = {"kind": kind, "url": url}
        if vid:
            item["vid"] = vid
        videos.append(item)

    for v in soup.find_all("video"):
        src = v.get("src") or v.get("data-src") or ""
        if src:
            add("mp4", src)
    for m in soup.find_all("mpvideo"):
        src = m.get("src") or m.get("data-src") or ""
        if src:
            add("mp4", src)
    for f in soup.find_all("iframe"):
        src = f.get("data-src") or f.get("src") or ""
        if "v.qq.com" in (src or ""):
            m = re.search(r"[?&]vid=([A-Za-z0-9]+)", src or "")
            vid = m.group(1) if m else ""
            if vid:
                add("qq_video", f"https://v.qq.com/x/page/{vid}.html", vid=vid)
            elif src:
                add("qq_video", src)
    return videos


def _download_file(url: str, path: Path, *, cred: dict[str, Any] | None, timeout: float) -> None:
    headers = {
        "User-Agent": USER_AGENT,
        "Referer": "https://mp.weixin.qq.com/",
        "Accept": "*/*",
    }
    cookies: dict[str, str] = {}
    if cred:
        pt = str(cred.get("pass_ticket") or "").strip()
        if pt:
            cookies["pass_ticket"] = _fully_unquote(pt)
    sess = requests.Session()
    sess.trust_env = False
    resp = sess.get(url, headers=headers, cookies=cookies, timeout=timeout, stream=True)
    try:
        resp.raise_for_status()
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("wb") as fh:
            for chunk in resp.iter_content(chunk_size=65536):
                if chunk:
                    fh.write(chunk)
    finally:
        close = getattr(resp, "close", None)
        if callable(close):
            close()


def download_article_videos(
    art: dict[str, Any],
    out_dir: Path | str,
    *,
    cred: dict[str, Any] | None = None,
    timeout: float = 30.0,
) -> list[dict[str, Any]]:
    """Download mp4 videos into ``out_dir/视频/``; v.qq.com links are kept as-is."""
    videos: list[dict[str, Any]] = []
    for v in art.get("videos") or []:
        if isinstance(v, dict):
            videos.append(dict(v))
    if not videos:
        return videos
    dl_dir = Path(out_dir) / VIDEO_DL_DIRNAME
    safe = re.sub(r'[\\/:*?"<>|]+', "_", str(art.get("title") or "article"))[:48] or "article"
    n = 0
    for v in videos:
        if v.get("kind") != "mp4":
            continue
        url = str(v.get("url") or "").strip()
        if not url:
            continue
        n += 1
        path = dl_dir / f"{safe}_{n:02d}.mp4"
        try:
            _download_file(url, path, cred=cred, timeout=timeout)
            v["local_path"] = str(path)
        except Exception as exc:  # noqa: BLE001
            v["error"] = describe_exception(exc)
    return videos


def article_to_html_document(art: dict[str, Any]) -> str:
    normalized = str(art.get("normalized_html") or "").strip()
    if normalized and "<!DOCTYPE html>" in normalized:
        return normalized

    title = html.escape(str(art.get("title") or "(无标题)"))
    link = html.escape(str(art.get("link") or ""))
    when = html.escape(str(art.get("publish_at") or ""))
    body_html = str(art.get("body_html") or "")
    if not body_html.strip():
        body_html = f"<pre>{html.escape(str(art.get('body_text') or ''))}</pre>"
    meta_bits = []
    if link:
        meta_bits.append(f'<p class="meta">来源：<a href="{link}">{link}</a></p>')
    if when:
        meta_bits.append(f'<p class="meta">发布时间：{when}</p>')
    meta = "\n".join(meta_bits)
    return (
        "<!DOCTYPE html>\n"
        '<html lang="zh-CN">\n'
        "<head>\n"
        '<meta charset="utf-8"/>\n'
        f"<title>{title}</title>\n"
        "<style>\n"
        "body{font-family:Microsoft YaHei,Segoe UI,sans-serif;max-width:760px;"
        "margin:32px auto;padding:0 20px;line-height:1.7;color:#1a1a1a;}\n"
        "h1{font-size:1.6rem;margin-bottom:0.4em;}\n"
        ".meta{color:#666;font-size:0.9rem;}\n"
        "img{max-width:100%;height:auto;}\n"
        "</style>\n"
        "</head>\n"
        "<body>\n"
        f"<h1>{title}</h1>\n"
        f"{meta}\n"
        f'<div class="content">{body_html}</div>\n'
        "</body>\n"
        "</html>\n"
    )


def _set_run_font(run: Any, *, size: float | None = None, bold: bool | None = None,
                  color: str | None = None, italic: bool | None = None) -> None:
    """Set run font incl. east-asian font name (works for Chinese text)."""
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Microsoft YaHei"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "微软雅黑")


def _docx_add_hyperlink(paragraph: Any, url: str, text: str) -> None:
    """Add a clickable hyperlink run to a paragraph (needs OOXML plumbing)."""
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    rpr.append(style)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    color.set(qn("w:themeColor"), "hyperlink")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "微软雅黑")
    rpr.append(rfonts)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _docx_inline(node: Any) -> list[tuple[str, dict[str, Any]]]:
    """Flatten an inline HTML subtree into [(text, style_overrides)]."""
    out: list[tuple[str, dict[str, Any]]] = []
    if isinstance(node, str):
        if node.strip():
            out.append((node, {}))
        return out
    if not isinstance(node, Tag):
        return out
    name = (node.name or "").lower()
    base: dict[str, Any] = {}
    if name in ("strong", "b"):
        base["bold"] = True
    elif name in ("em", "i"):
        base["italic"] = True
    elif name == "a":
        base["link"] = str(node.get("href") or "").strip()
    elif name == "br":
        out.append(("\n", {}))
        return out
    elif name in ("p", "div", "section", "span", "h1", "h2", "h3", "h4", "h5", "h6"):
        pass
    else:
        pass
    for child in node.children:
        for text, extra in _docx_inline(child):
            merged = dict(base)
            merged.update(extra)
            out.append((text, merged))
    return out


def _docx_download_image(url: str, timeout: float = 10.0) -> bytes | None:
    """Download an article image, bypassing the local capture proxy."""
    if not url or not url.startswith(("http://", "https://")):
        return None
    try:
        sess = requests.Session()
        sess.trust_env = False
        resp = sess.get(
            url,
            headers={
                "User-Agent": USER_AGENT,
                "Referer": "https://mp.weixin.qq.com/",
            },
            timeout=timeout,
        )
        resp.raise_for_status()
        data = resp.content
        if not data or len(data) < 64 or len(data) > 20 * 1024 * 1024:
            return None
        return data
    except Exception:
        return None


def _render_docx_body(doc: Document, body_html: str) -> None:
    """Append styled paragraphs / headings / lists / images from body HTML."""
    soup = BeautifulSoup(body_html or "", "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()

    def add_image(img: Tag) -> None:
        src = str(img.get("src") or img.get("data-src") or "").strip()
        data = _docx_download_image(src) if src else None
        if data:
            try:
                from io import BytesIO

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(BytesIO(data), width=Cm(14))
                return
            except Exception:
                pass
        p = doc.add_paragraph()
        run = p.add_run("[图片]")
        _set_run_font(run, size=10, color="888888")

    def walk(node: Any, list_level: int = 0) -> None:
        if not isinstance(node, Tag):
            return
        name = (node.name or "").lower()

        if name == "img":
            add_image(node)
            return

        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(name[1])
            text = "".join(t for t, _ in _docx_inline(node)).strip()
            h = doc.add_heading(text or "(无标题)", level=min(level, 3))
            for run in h.runs:
                _set_run_font(run, size={1: 18, 2: 15, 3: 13}.get(level, 12), bold=True,
                              color="1A1A1A")
            return

        if name in ("ul", "ol"):
            for idx, li in enumerate(node.find_all("li", recursive=False), start=1):
                parts = _docx_inline(li)
                text = "".join(t for t, _ in parts).strip()
                bullet = "• " if name == "ul" else f"{idx}. "
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(bullet)
                _set_run_font(run, size=11, bold=True)
                for txt, over in parts:
                    if not txt.strip():
                        continue
                    if over.get("link"):
                        _docx_add_hyperlink(p, over["link"], txt)
                    else:
                        run = p.add_run(txt)
                        _set_run_font(run, size=11, bold=over.get("bold"),
                                      italic=over.get("italic"))
            return

        if name == "blockquote":
            text = "".join(t for t, _ in _docx_inline(node)).strip()
            if text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text)
                _set_run_font(run, size=11, italic=True, color="555555")
            return

        if name == "p":
            # 段落可能同时包含文本与图片
            imgs = node.find_all("img", recursive=False)
            parts = _docx_inline(node)
            has_text = any(t.strip() for t, _ in parts)
            if not has_text and not imgs:
                return
            if imgs and has_text:
                for img in imgs:
                    img.decompose()
                parts = _docx_inline(node)
                has_text = any(t.strip() for t, _ in parts)
            if has_text:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.6
                p.paragraph_format.space_after = Pt(6)
                for txt, over in parts:
                    if not txt.strip():
                        continue
                    if over.get("link"):
                        _docx_add_hyperlink(p, over["link"], txt)
                    else:
                        run = p.add_run(txt)
                        _set_run_font(run, size=11, bold=over.get("bold"),
                                      italic=over.get("italic"))
            for img in imgs:
                add_image(img)
            return

        if name in ("div", "section", "article", "span"):
            for child in node.children:
                walk(child, list_level)
            return

        # fallback: any other block-ish tag
        for child in node.children:
            walk(child, list_level)

    root = soup.find(["div", "section", "article", "body"]) or soup
    walk(root)


def article_to_docx(art: dict[str, Any]) -> bytes:
    """Build a styled .docx from article title / meta / rich body HTML."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "微软雅黑")

    title = str(art.get("title") or "(无标题)").strip()
    link = str(art.get("link") or "").strip()
    when = str(art.get("publish_at") or "").strip()
    body_html = str(art.get("body_html") or "").strip()
    body_text = str(art.get("body_text") or "").strip()

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        _set_run_font(run, size=18, bold=True, color="1A1A1A")

    if when:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(when)
        _set_run_font(run, size=9, color="888888")

    if link:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_add_hyperlink(p, link, link)
        for run in p.runs:
            _set_run_font(run, size=9, color="0563C1")

    doc.add_paragraph()

    if body_html:
        _render_docx_body(doc, body_html)
    elif body_text:
        for block in re.split(r"\n\s*\n", body_text):
            line = block.strip()
            if line:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.6
                run = p.add_run(line)
                _set_run_font(run, size=11)
    else:
        doc.add_paragraph("（无正文）")

    videos = [
        v
        for v in (art.get("videos") or [])
        if isinstance(v, dict) and (v.get("local_path") or v.get("url"))
    ]
    if videos:
        p = doc.add_paragraph()
        run = p.add_run("视频：")
        _set_run_font(run, size=11, bold=True)
        for v in videos:
            target = str(v.get("local_path") or v.get("url") or "").strip()
            if not target:
                continue
            p2 = doc.add_paragraph()
            _docx_add_hyperlink(p2, str(v.get("url") or target), target)
            for run in p2.runs:
                _set_run_font(run, size=10, color="0563C1")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _normalize_fmt(fmt: str) -> str:
    key = (fmt or "markdown").lower().strip()
    aliases = {
        "md": "markdown",
        "markdown": "markdown",
        "html": "html",
        "txt": "txt",
        "text": "txt",
        "json": "json",
        "csv": "csv",
        "word": "word",
        "docx": "word",
        "doc": "word",
    }
    return aliases.get(key, key)


def render_article_export(art: dict[str, Any], fmt: str) -> str:
    """Return text payload. For Word use ``article_to_docx`` / ``write_article_export``."""
    key = _normalize_fmt(fmt)
    if key == "html":
        return article_to_html_document(art)
    if key == "markdown":
        return article_to_markdown(art)
    if key == "txt":
        return article_to_txt(art)
    if key == "json":
        return article_to_json(art)
    if key == "csv":
        buf = io.StringIO()
        writer = csv.DictWriter(buf, fieldnames=list(CSV_COLUMNS))
        writer.writeheader()
        writer.writerow(article_to_csv_row(art))
        return buf.getvalue()
    if key == "word":
        raise ValueError("Word 为二进制格式，请使用 write_article_export / article_to_docx")
    raise ValueError(f"不支持的导出格式: {fmt}")


def write_article_export(path: Path | str, art: dict[str, Any], fmt: str) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    key = _normalize_fmt(fmt)
    if key == "word":
        path.write_bytes(article_to_docx(art))
    elif key == "csv":
        _write_csv(path, [article_to_csv_row(art)])
    else:
        path.write_text(render_article_export(art, key), encoding="utf-8")
    return path


def safe_export_filename(title: str, *, ext: str, index: int = 0) -> str:
    safe = re.sub(r'[\\/:*?"<>|]+', "_", (title or "article").strip())[:48] or "article"
    if index > 0:
        return f"{index:02d}_{safe}.{ext}"
    return f"{safe}.{ext}"


def _fetch_with_retry(
    fetch: Callable[..., dict[str, Any]],
    link: str,
    *,
    cred: dict[str, Any] | None,
    retries: int = 2,
    retry_delay_s: float = 1.0,
) -> dict[str, Any]:
    """Fetch article with transient-network retries (timeout / connection)."""
    last: Exception | None = None
    for attempt in range(retries + 1):
        try:
            return fetch(link, cred=cred)
        except (requests.Timeout, requests.ConnectionError) as exc:
            last = exc
            if attempt < retries:
                time.sleep(retry_delay_s * (attempt + 1))
    if last is None:  # pragma: no cover - unreachable
        last = requests.RequestException("请求失败")
    raise last


def batch_export_articles(
    articles: list[dict[str, Any]],
    *,
    out_dir: Path | str,
    fmt: str = "markdown",
    fetch_article: Callable[..., dict[str, Any]] | None = None,
    cred: dict[str, Any] | None = None,
    on_progress: Callable[[str], None] | None = None,
    sleep_s: float = 0.3,
    csv_path: Path | str | None = None,
    download_videos: bool = True,
) -> dict[str, Any]:
    """Fetch each article body and write into ``out_dir``.

    ``fetch_article(url, cred=...)`` defaults to ``fetch_and_parse_article``.
    CSV 格式：所有文章合并进**一个** csv（``csv_path`` 或 ``out_dir/导出文章.csv``）。
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    key = _normalize_fmt(fmt)
    ext = extension_for_article_format(key)
    fetch = fetch_article or fetch_and_parse_article
    ok_n = 0
    failed_n = 0
    errors: list[str] = []
    written: list[str] = []
    rows_out: list[dict[str, str]] = []

    for i, row in enumerate(articles, start=1):
        link = str(row.get("link") or "").strip()
        title = str(row.get("title") or f"article_{i}").strip()
        if on_progress:
            on_progress(f"正在导出 {i}/{len(articles)}：{title[:28]}")
        if not link:
            failed_n += 1
            errors.append(f"{title}: 无链接")
            continue
        try:
            parsed = _fetch_with_retry(fetch, link, cred=cred)
            if not parsed.get("publish_at") and row.get("publish_at"):
                parsed["publish_at"] = row.get("publish_at")
            if not parsed.get("publish_ts") and row.get("publish_ts"):
                parsed["publish_ts"] = row.get("publish_ts")
            if not parsed.get("title") or parsed.get("title") == "(无标题)":
                parsed["title"] = title or parsed.get("title")
            if download_videos:
                parsed["videos"] = download_article_videos(parsed, out_dir, cred=cred)
            if key == "csv":
                rows_out.append(article_to_csv_row(parsed))
            else:
                fname = safe_export_filename(str(parsed.get("title") or title), ext=ext, index=i)
                path = write_article_export(out_dir / fname, parsed, key)
                written.append(str(path))
            ok_n += 1
        except Exception as exc:  # noqa: BLE001
            failed_n += 1
            errors.append(f"{title}: {describe_exception(exc)}")
        if sleep_s > 0 and i < len(articles):
            time.sleep(sleep_s)

    if key == "csv":
        final_csv = Path(csv_path) if csv_path else out_dir / "导出文章.csv"
        _write_csv(final_csv, rows_out)
        written = [str(final_csv)]

    return {
        "ok": ok_n,
        "failed": failed_n,
        "errors": errors,
        "written": written,
        "out_dir": str(out_dir),
        "out_file": written[0] if written else "",
        "fmt": key,
    }


def format_key_for_article_label(label: str) -> str:
    for key, lb in ARTICLE_EXPORT_FORMATS.items():
        if lb == label:
            return key
    return "markdown"


def extension_for_article_format(fmt: str) -> str:
    key = _normalize_fmt(fmt)
    if key == "markdown":
        return "md"
    if key == "html":
        return "html"
    if key == "txt":
        return "txt"
    if key == "json":
        return "json"
    if key == "csv":
        return "csv"
    if key == "word":
        return "docx"
    return "md"


def fetch_article_html(
    url: str,
    *,
    cred: dict[str, Any] | None = None,
    timeout: float = 25.0,
    session: requests.Session | None = None,
    retries: int = 2,
    retry_delay_s: float = 1.0,
) -> str:
    """Fetch article page HTML (direct to WeChat, bypass system proxy)."""
    url = (url or "").strip()
    if not url:
        raise ValueError("文章链接为空")

    headers = {
        "User-Agent": USER_AGENT,
        "Referer": "https://mp.weixin.qq.com/",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    }
    cookies: dict[str, str] = {}
    if cred:
        pt = str(cred.get("pass_ticket") or "").strip()
        uin = str(cred.get("uin") or "").strip()
        if pt:
            cookies["pass_ticket"] = _fully_unquote(pt)
        if uin:
            cookies["wxuin"] = _fully_unquote(uin)

    sess = session or requests.Session()
    sess.trust_env = False
    last_exc: Exception | None = None
    for attempt in range(retries + 1):
        try:
            resp = sess.get(url, headers=headers, cookies=cookies, timeout=timeout)
            resp.raise_for_status()
            resp.encoding = resp.apparent_encoding or resp.encoding or "utf-8"
            return resp.text
        except (requests.Timeout, requests.ConnectionError) as exc:
            last_exc = exc
            if attempt < retries:
                time.sleep(retry_delay_s * (attempt + 1))
        except requests.exceptions.HTTPError as exc:
            status = exc.response.status_code if exc.response is not None else 0
            if status in (429, 500, 502, 503, 504) and attempt < retries:
                time.sleep(retry_delay_s * (attempt + 1))
                continue
            raise
    if last_exc is None:  # pragma: no cover - unreachable
        last_exc = requests.RequestException("请求失败")
    raise last_exc


def fetch_and_parse_article(
    url: str,
    *,
    cred: dict[str, Any] | None = None,
    timeout: float = 25.0,
) -> dict[str, Any]:
    html_text = fetch_article_html(url, cred=cred, timeout=timeout)
    return parse_wechat_article_html(html_text, source_url=url)


_BIZ_URL_RE = re.compile(r"__biz=([A-Za-z0-9_%\-]+)")
_BIZ_URL_RE = re.compile(r"__biz=([A-Za-z0-9_%\-]+)")
_BIZ_SCRIPT_RE = re.compile(r'var\s+biz\s*=\s*["\']([^"\']+)["\']')
_BIZ_JSON_RE = re.compile(r'"__biz"\s*:\s*"([^"]+)"')

def extract_biz_from_html(html: str) -> str:
    """从公众号文章 HTML 提取 __biz（og:url / var biz / JSON 字段）。

    短链（/s/xxx 无查询参数）URL 里没有 __biz，但文章页 HTML 里有。
    """
    if not html:
        return ""
    for pat in (_BIZ_URL_RE, _BIZ_SCRIPT_RE, _BIZ_JSON_RE):
        m = pat.search(html)
        if m:
            try:
                return unquote(m.group(1))
            except Exception:
                return m.group(1)
    return ""


def fetch_biz_from_url(url: str, *, timeout: float = 15.0, session: Any | None = None) -> str:
    """直接抓取公开文章页并提取 __biz；失败返回空串（无需凭证）。"""
    try:
        html = fetch_article_html(url, timeout=timeout, session=session)
        return extract_biz_from_html(html)
    except Exception:
        return ""
